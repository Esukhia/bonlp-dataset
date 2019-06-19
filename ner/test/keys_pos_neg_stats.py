import argparse                                                                               
import os
import glob

from docx import Document                                                                     
from tqdm import tqdm                                                                         
                                                                                              
                                                                                              
def find_stats(fn):                                                                           
    with open(fn, 'rb') as f:                                                                 
        doc = Document(f)
    print(fn)    
    print(doc.tables[6].rows[0].cells[1].text)
    #for table in doc.tables:
    #    print(table)

def stats_to_csv(stats):
    pass
        
        
        
if __name__ == "__main__":
    ap = argparse.ArgumentParser(add_help=False)
    ap.add_argument("--path", type=str, help="path to corpus")
    args = ap.parse_args()

    fns = sorted(glob.glob(os.path.join(args.path, "*")))

    for fn in fns:
        find_stats(fn)
